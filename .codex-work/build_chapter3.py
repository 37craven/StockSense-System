from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\User\source\repos\StockSense-System\output\documents\StockSense_Chapter_3_Technical_Background.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "1F4D78"
BLUE = "2E74B5"
LIGHT = "F4F6F9"
GRID = "B8C2CC"
MUTED = "666666"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.33

for name, size, color, before, after in [
    ("Title", 26, NAVY, 0, 8), ("Subtitle", 13, MUTED, 0, 18),
    ("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4)]:
    st = styles[name]
    st.font.name = "Calibri"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Diagram" not in styles:
    st = styles.add_style("Diagram", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Consolas"; st.font.size = Pt(8.5)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Consolas"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    st.paragraph_format.space_before = Pt(8); st.paragraph_format.space_after = Pt(8); st.paragraph_format.line_spacing = 1.0

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        x = tcMar.find(qn(f"w:{m}"))
        if x is None: x = OxmlElement(f"w:{m}"); tcMar.append(x)
        x.set(qn("w:w"), str(v)); x.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); e = OxmlElement("w:tblHeader"); e.set(qn("w:val"), "true"); trPr.append(e)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); c.text=h; shade(c, LIGHT); cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: set_font(r, size=9.5, bold=True, color=NAVY)
    set_repeat_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i, val in enumerate(row):
            cells[i].width=Inches(widths[i]); cells[i].text=str(val); cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.05
                for r in p.runs: set_font(r, size=9)
    # exact fixed geometry
    tblPr=t._tbl.tblPr; tblW=tblPr.first_child_found_in("w:tblW"); tblW.set(qn("w:w"),"9360"); tblW.set(qn("w:type"),"dxa")
    for old_layout in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old_layout)
    layout=OxmlElement("w:tblLayout"); layout.set(qn("w:type"),"fixed"); tblPr.append(layout)
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); tblPr.append(ind)
    grid=t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"),str(round(w*1440))); grid.append(gc)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            tcW=c._tc.get_or_add_tcPr().first_child_found_in("w:tcW"); tcW.set(qn("w:w"),str(round(widths[i]*1440))); tcW.set(qn("w:type"),"dxa")
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def h(text, level=1): doc.add_heading(text, level=level)
def p(text, bold_lead=None):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r=para.add_run(bold_lead); set_font(r,bold=True); r=para.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r=para.add_run(text); set_font(r)
    return para

def page_field(paragraph):
    run=paragraph.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); run._r.addnext(fld)

# Running header/footer
hp=sec.header.paragraphs[0]; hp.text="STOCKSENSE  |  CHAPTER 3 - TECHNICAL BACKGROUND"; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: set_font(r,size=8.5,bold=True,color=MUTED)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=fp.add_run("Page "); set_font(r,size=9,color=MUTED); page_field(fp)

# Title page
doc.add_paragraph().paragraph_format.space_after=Pt(72)
tp=doc.add_paragraph(style="Title"); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(tp.add_run("CHAPTER 3"),size=26,bold=True,color=NAVY)
sp=doc.add_paragraph(style="Subtitle"); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(sp.add_run("TECHNICAL BACKGROUND OF STOCKSENSE"),size=15,bold=True,color=BLUE)
p1=doc.add_paragraph(); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p1.add_run("A Web-Based Inventory, Sales, Service, and Motorcycle Build Management System with AI-Assisted Support"),size=12,color=MUTED)
doc.add_paragraph().paragraph_format.space_after=Pt(54)
meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(meta.add_run("Implementation-based manuscript draft\nPrepared from the StockSense .NET 8 solution and Gemini chatbot service\nAugust 2026"),size=10.5,color=MUTED)
doc.add_page_break()

h("Scope and Accuracy Note",1)
p("This chapter describes the technology and behavior that can be verified in the current StockSense source code. Recommended hardware and network values are presented as deployment requirements because the repository does not record the developers' physical machines or the final production service plan. Features are stated with their actual boundaries: reorder automation generates draft order slips that staff can confirm and email to suppliers; barcode reading is implemented through a browser camera; online payment is recorded but no payment gateway is integrated; and reporting currently consists of operational dashboards, histories, and generated PDF documents.")

h("1. Details of the Technology to Be Used",1)
h("1.1 Software Requirements",2)
p("StockSense uses a layered .NET architecture. The browser interface is implemented with Blazor, the server exposes ASP.NET Core controllers and application services, and Entity Framework Core maps domain objects to SQL Server. A separate Python service provides AI-assisted support using Google Gemini. Table 1 summarizes the software actually declared or used by the source projects.")
software = [
("IDE / tooling","Visual Studio 2026","2026","Primary environment for editing, building, debugging, database migration, and testing. It was selected for its integrated support for C#, Razor, ASP.NET Core, Blazor, Git, and .NET tooling."),
("Runtime",".NET SDK","8.0","Builds and runs every StockSense project. Long-term-support .NET 8 provides ASP.NET Core, Blazor, dependency injection, security middleware, and cross-platform deployment."),
("Language","C#","12 / .NET 8 compiler","Implements domain entities, services, repositories, controllers, and tests with static typing, asynchronous APIs, and strong tooling."),
("Language","Razor / HTML / CSS / JavaScript","Blazor .NET 8 / browser standards","Builds interactive pages and camera-based barcode scanning. Razor combines HTML markup with reusable C# components."),
("Frontend framework","Blazor Web App and WebAssembly","8.0.23","Provides interactive server and browser-rendered components while allowing the client to call same-origin API endpoints."),
("UI library","BlazorBlueprint","Components/Primitives 3.14.1; Lucide icons 2.0.0","Supplies accessible, reusable controls and a consistent Tailwind-based visual system without a custom Tailwind build."),
("Backend framework","ASP.NET Core Web API","8.0","Hosts controllers, authentication, authorization, rate limiting, antiforgery middleware, static assets, and server-side services."),
("Architecture","Clean Architecture-inspired layers","Domain, Application, Infrastructure, Web, Client","Separates core records and rules from use cases, database access, external services, and presentation, improving maintainability and testability."),
("Database","Microsoft SQL Server / Azure SQL-compatible SQL","EF Core provider 8.0.23","Stores identity, products, inventory metrics, suppliers, orders, sales, appointments, builds, services, and motorcycle compatibility. Local development uses LocalDB; production can use Azure SQL through a connection string."),
("ORM","Entity Framework Core","8.0.23","Maps C# entities to relational tables, applies migrations, parameterizes normal queries, and supports transactions and optimistic concurrency."),
("Identity/security","ASP.NET Core Identity and BCrypt.Net-Next","Identity EF 8.0.23; BCrypt 4.1.0","Manages accounts, cookies, roles, confirmation, lockout, and password verification. The application pre-hashes with SHA-256 and then uses BCrypt cost 12."),
("Documents","QuestPDF","2026.7.2","Generates order-slip and barcode-label PDF documents with controlled layout."),
("Barcode imaging","ZXing.Net / SkiaSharp binding","0.16.11 / 0.16.22","Generates and reads EAN-13 and QR imagery; browser JavaScript supplies camera frames for POS scanning."),
("Image processing","SixLabors.ImageSharp","3.1.11","Validates and processes uploaded product images on the server."),
("Email","MailKit / MimeKit","4.15.0 / 4.15.1","Sends confirmation, recovery, quotation, and order-slip email through configured SMTP using MIME messages and attachments."),
("AI microservice","Python + FastAPI + Uvicorn","FastAPI >=0.115,<1; Uvicorn >=0.34,<1","Runs the independent chatbot HTTP service and validates request/response contracts."),
("AI SDK / model","Google GenAI SDK / Gemini","google-genai >=1,<2; default gemini-3.1-flash-lite","Generates role-appropriate natural-language answers after the chatbot retrieves permitted shop data."),
("AI data/search","pyodbc, ddgs, python-dotenv","pyodbc >=5.2,<6; ddgs >=9,<10; dotenv >=1,<2","Connects the chatbot to SQL Server through ODBC Driver 18, supports bounded web lookup, and loads local environment configuration."),
("Testing","xUnit, EF Core InMemory, coverlet","xUnit 2.5.3; test SDK 17.8; coverlet 6.0","Exercises inventory mathematics, workflows, authorization, chatbot integration, compatibility, and UI helper behavior."),
("Version control","Git and GitHub","Current supported release","Tracks source changes, supports branches and collaboration, and provides an auditable development history.")]
table(["Classification","Software","Version","Purpose, selection, and relevant features"],software,[1.0,1.45,1.15,2.9])

h("1.2 Hardware Requirements",2)
p("Table 2 gives practical minimum and recommended specifications for development and demonstration. They are requirements inferred from the workload, not a record of a specific developer laptop. The server and chatbot may share one machine during development but should be independently deployable in production.")
hardware=[
("Processor","Minimum: 4-core 64-bit CPU; recommended: modern 6- to 8-core Intel Core i5/i7, AMD Ryzen 5/7, or equivalent","Parallel compilation, SQL Server, browser sessions, tests, and the Python service benefit from multiple cores. No local GPU is required because Gemini inference runs in Google's service."),
("Memory (RAM)","Minimum: 8 GB; recommended: 16 GB or more","Allows Visual Studio, the .NET host, SQL Server/LocalDB, browser tools, and chatbot service to run together without excessive paging."),
("Storage","Minimum: 10 GB free SSD; recommended: 256 GB or larger SSD with 20 GB free for tools and data","An SSD shortens builds, package restoration, database migration, image handling, and test execution. Backups need additional protected storage."),
("Display","1366x768 minimum; 1920x1080 recommended","Supports responsive UI review, IDE panels, debugging, and side-by-side browser testing."),
("Network","100 Mbps Ethernet or stable dual-band Wi-Fi; internet access","Required for package restoration, GitHub, SMTP, cloud SQL, Gemini, deployment, and remote browser testing."),
("Operating system","Windows 10/11 recommended for Visual Studio and LocalDB; supported Linux/Windows server for deployment","The .NET application is cross-platform, while the documented local database workflow and ODBC setup are simplest on Windows.")]
table(["Hardware component","Technical specification","Rationale and development support"],hardware,[1.2,2.2,3.1])

h("1.3 Network Requirements",2)
p("Development requires a stable internet connection of at least 10 Mbps, with 25 Mbps or higher recommended for package restoration, source synchronization, cloud deployment, and video-capable remote testing. Latency should remain reasonably low and packet loss minimal because the AI module calls an external Gemini endpoint and waits for a generated response. Local-only work can continue for core CRUD features when LocalDB is used, but AI, SMTP, and cloud resources will be unavailable.")
p("On a local area network, developer devices should use DHCP or reserved addresses and permit HTTPS traffic to the ASP.NET Core host. If a phone is used for camera scanning, it must reach the test server and the browser must be granted camera permission; secure contexts are normally required outside localhost. Production traffic should use HTTPS on port 443. Database port 1433 and ODBC connectivity should be restricted to trusted application hosts rather than exposed to public clients.")
p("Cloud deployment needs outbound HTTPS access to Google Gemini, access to the selected SQL Server or Azure SQL endpoint, DNS resolution, and SMTP connectivity to the configured mail server. Secrets such as database, SMTP, and Gemini credentials must be supplied through environment variables, user secrets, or a managed secret store. The chatbot should be private or protected by service-to-service authentication because its direct FastAPI endpoint currently trusts the role supplied in JSON; public exposure would permit a caller to claim a staff role.")
p("Testing should cover interrupted requests, slow AI responses, unavailable email, database retry behavior, concurrent stock updates, camera permission denial, and different desktop/mobile browsers. Firewalls and reverse proxies must preserve the application origin and forwarded HTTPS information so secure cookie behavior remains correct.")

h("1.4 Peopleware Requirements",2)
people=[
("Project Manager / Systems Analyst","Defines scope, schedules work, coordinates stakeholders, translates shop workflows into requirements, controls risks and changes, and ensures manuscript and implementation remain aligned."),
("Frontend Developer","Builds accessible Blazor pages, responsive layouts, forms, dashboards, role-specific navigation, camera scanning, and clear validation/error feedback."),
("Backend Developer","Implements ASP.NET Core controllers, application services, transactions, authorization policies, external integrations, and API contracts."),
("Database Administrator / Data Engineer","Designs and reviews SQL Server schemas, migrations, indexes, backups, least-privilege accounts, retention, recovery, and production connection security."),
("AI / Integration Developer","Maintains the FastAPI/Gemini service, retrieval tools, role-aware prompting, ASP.NET proxy, timeouts, safe data exposure, and service authentication."),
("QA / Test Engineer","Creates unit, integration, authorization, concurrency, mobile, barcode, usability, and failure-mode tests; records defects and verifies corrections."),
("UI/UX Designer","Models customer and staff journeys, checks terminology with shop users, specifies responsive behavior, and validates accessibility and task efficiency."),
("DevOps / Cloud Administrator","Configures application hosting, HTTPS, DNS, Azure SQL or SQL Server, monitoring, backups, secret injection, deployment approvals, and incident recovery."),
("Motorcycle Shop Domain Expert","Validates product, supplier, compatibility, safety-stock, service, appointment, and work-order rules so technical behavior reflects actual operations."),
("End Users: Admin, Employee, Customer","Perform acceptance testing and provide feedback on management controls, POS/service workflows, booking/build submission, and AI assistance.")]
table(["Role","Development responsibilities"],people,[1.8,4.7])

h("2. Project Technical Description",1)
h("2.1 Overall System Architecture",2)
p("StockSense is a browser-based system for a motorcycle parts and service shop. The visible screens are Blazor components. Some are rendered interactively by the server and others can execute as WebAssembly in the browser. Both use ASP.NET Core endpoints in the StockSense Web host. Behind those endpoints, application and infrastructure services apply business rules, talk to SQL Server through Entity Framework Core, generate documents, and send email.")
diagram = """+-------------------- USERS --------------------+\n| Customer browser | Employee POS | Admin panel |\n+-------------------------+----------------------+\n                          | HTTPS / Blazor / JSON\n+-------------------------v----------------------+\n| StockSense.Web (.NET 8 Blazor + ASP.NET Core) |\n| Identity cookies | RBAC | Controllers | UI     |\n+----------+----------------------+---------------+\n           | application services | authorized AI proxy\n+----------v-----------+     +----v----------------------+\n| Domain/Application   |     | FastAPI Chatbot          |\n| inventory, orders,   |     | role-aware tools +       |\n| sales, service rules |     | Google Gemini synthesis  |\n+----------+-----------+     +----+----------------------+\n           | EF Core              | permitted read queries\n+----------v----------------------v----------------------+\n| SQL Server / Azure SQL-compatible relational data     |\n+-------------------------------------------------------+\nOther integrations: SMTP email | QuestPDF | camera/EAN-13 scanning"""
dp=doc.add_paragraph(style="Diagram"); dp.add_run(diagram)
p("When a person submits a form, the browser sends an HTTPS request to a controller. The controller checks the signed-in identity and, where configured, the required role. It validates or maps the request, calls the relevant service or repository, and returns JSON or an updated screen. Entity Framework Core translates normal database operations into parameterized SQL. Important stock-changing workflows use serializable transactions or row-version checks so simultaneous users do not silently overwrite each other.")
p("AI-assisted support follows an additional path. The browser calls the authenticated StockSense endpoint at /api/assistance. The ASP.NET server derives the user's highest role from trusted identity claims and forwards the message, limited conversation history, and role to the FastAPI chatbot. The chatbot chooses only the data tools allowed for that role, obtains current shop facts from SQL Server where required, and asks Gemini to form a clear answer. The reply returns through ASP.NET to the browser. This proxy prevents the normal web client from choosing its own role, although the FastAPI service itself still needs service authentication or private networking before public production use.")
p("For deployment, the .NET host can run on a Windows or Linux web service behind HTTPS, with static/client assets served from the same application. SQL Server may be local during development and Azure SQL-compatible in production. The chatbot is independently deployable as a Python web service. The repository contains configuration points for these services but no complete infrastructure-as-code definition, so the final hosting region, scale tier, backup policy, and availability target must be decided by the deployment team.")

h("2.2 Special Features of the System",2)
h("2.2.1 AI-Assisted Support",3)
p("The assistant is not a free-form chatbot attached directly to the browser. StockSense authenticates the user first and passes a server-derived role to a separate FastAPI service. Customer tools are limited to public products, active services, pre-built packages, appointment-slot snapshots, and compatibility guidance. Employees can obtain operational inventory, safety-stock, sales, order-slip, and work-order information. Administrators receive those capabilities plus aggregate role counts. Database results are capped and the assistant is instructed not to invent product fitment or shop facts.")
p("The service uses the Google GenAI SDK and defaults to the Gemini 3.1 Flash Lite model, configurable through an environment variable. Requests allow a message of up to 8,000 characters and a bounded twelve-turn history. Tool execution is also bounded. These limits reduce cost, latency, and abuse. The chatbot can return navigation guidance and current information, but it does not replace authoritative booking or transaction operations inside StockSense.")
h("2.2.2 Safety Stock and Reorder Point Logic",3)
p("Safety stock is the extra quantity kept to absorb uncertain demand or supplier delay. Reorder point is the inventory level at which replenishment should begin. StockSense estimates daily demand from sale quantities plus recorded lost-sales quantities, and builds a complete daily series that includes zero-sale days. In simplified form, lead-time demand equals average daily demand multiplied by expected lead time. Safety stock combines demand variability, lead-time behavior, and the chosen service-level factor. The reorder point is lead-time demand plus safety stock. Minimum, maximum, package-size, and minimum-order constraints are then applied.")
p("The system calculates inventory position as current stock plus incoming stock minus reserved stock. When that position reaches the policy threshold, it recommends a quantity, rounds it to package or minimum-order rules, respects maximum stock, and generates a draft order slip grouped by supplier. It deliberately avoids generating a second draft when an open order already covers the product. A staff member reviews the draft and confirms it. Once confirmed, StockSense can send the order slip to the supplier's registered email address. The supplier still fulfills the order outside the system, and staff record receiving when the items arrive.")
h("2.2.3 Staged Cold-Start Policy",3)
p("A new product has too little sales history for a stable statistical estimate. StockSense therefore uses stages. With fewer than 30 usable days, Cold Start converts an initial weekly estimate to a daily value and uses a configured buffer. From 30 to 59 days, Learning blends observed and initial demand equally. From 60 to 89 days, observed demand receives 70 percent weight. At 90 days or more, the Data-Driven stage relies on observed behavior. Observed supplier lead-time variability is used only after at least five completed orders. This staged policy prevents a few early sales from producing an exaggerated or misleading reorder level.")
h("2.2.4 Barcode and Mobile Scanning",3)
p("Products may receive a deterministic internal EAN-13 barcode beginning with the prefix 20, followed by the product identifier and a calculated check digit. The system can create PNG barcode/QR images and A6 PDF labels. In the employee POS page, browser JavaScript requests the phone or computer camera, detects a barcode, and asks the product repository for the matching item. This provides mobile-friendly scanning without requiring a separate physical scanner.")
h("2.2.5 Role-Based Access Control",3)
p("ASP.NET Core Identity provides accounts and the Admin, Employee, and Customer roles. Policies distinguish inventory staff from inventory administrators, and key sales, inventory, work-order, and account administration operations enforce roles. The AI proxy derives the role from server claims rather than accepting a role selected by the browser.")
p("RBAC should be described as implemented but still under hardening. In the current controllers, mechanics, services, and pre-built package mutations require authentication but do not consistently require an employee or administrator role. Some product responses can expose internal cost fields to authenticated customers. These endpoints should be corrected and regression-tested before claiming complete least-privilege enforcement.")
h("2.2.6 Real-Time Inventory Updates",3)
p("'Real time' means the database is updated as part of the completed business transaction, not that every open browser receives a push notification. A POS sale reloads authoritative product records inside a serializable transaction, deducts stock, saves the sale and line items, commits, and recalculates safety-stock metrics. Completing an appointment or build uses an atomic checkout service that creates the linked sale, deducts parts, marks the work completed, and then recalculates inventory policy. Row versions and transaction isolation reduce lost updates during simultaneous work.")

h("2.3 Web Application Modules",2)
modules=[
("2.3.1 User Access and Role Management Module","This module controls registration, confirmed-account sign-in, sign-out, password recovery, account management, and role assignment. ASP.NET Core Identity stores account data and issues a short-lived sliding authentication cookie. Administrators manage user roles, while navigation and endpoints adapt to Customer, Employee, or Admin responsibilities. Lockout and rate limits slow repeated login attempts. Two-factor and recovery management pages are present in the account area."),
("2.3.2 Inventory and Procurement Management Module","This module maintains products, suppliers, stock quantities, reserved and incoming quantities, barcodes, inventory settings, calculated metrics, and order slips. Administrators configure safety-stock policy; authorized staff view or recalculate it. The system calculates reorder recommendations and groups them into draft order slips by supplier. After staff review and confirmation, the confirmed order slip can be sent to the supplier's registered email address. The module also tracks order status and receiving and supports CRUD operations, low-stock visibility, image validation, concurrency control, and supplier/mobile details."),
("2.3.3 Sales and Transaction Management Module","The employee POS loads current products, builds a cart, accepts a recorded payment method/reference, scans barcodes through the device camera, completes a sale, deducts inventory, and presents a Sales Transaction Summary. Transaction history is filterable and limited to recent records. Sales use serializable database work to protect stock integrity. No external payment processor is integrated; payment data records what occurred outside or at the counter rather than charging a card or wallet."),
("2.3.4 Appointment and Service Management Module","Customers browse services and request appointments under their own server-derived identity. The system checks schedule conflicts and allows staff to assign mechanics, update status, and associate products. Confirmed parts may be reserved. On completion, the work-order checkout records a sale, deducts the parts, links the transaction, and closes the appointment atomically. Customers can view their own appointments and cancel an eligible pending booking."),
("2.3.5 Motorcycle Build Management Module","Customers select a motorcycle and compatible products or pre-built packages, submit a build request, and follow its status. Compatibility records link products, motorcycles, engine displacement ranges, and packages. Employees review requests and can connect installation work to an appointment. Completion follows the same protected work-order checkout path, converting selected parts into a linked transaction and inventory deduction."),
("2.3.6 Dashboard and Reporting Module","The dashboard converts current records into concise operational indicators: product count, low-stock count, current retail stock value, pending order information, and the most critical low-stock products. Transaction and order histories support day-to-day review, while QuestPDF produces order slips and barcode documents. This is an operational reporting foundation; broad exports, scheduled management reports, forecasting dashboards, and external BI integration are not demonstrated in the present code."),
("2.3.7 AI-Assisted Support Module","Customer and staff assistance pages send authenticated queries through the StockSense server to FastAPI and Gemini. Role-specific tool routing limits which database facts may be retrieved. The module handles FAQs, public product/service/package inquiries, appointment availability snapshots, compatibility guidance, operational status questions, and navigation assistance. It returns explanatory text but leaves actual bookings, stock adjustments, account changes, and sales to normal validated application endpoints.")]
for title,body in modules: h(title,3); p(body)

h("2.4 Special Hardware and Software Integration",2)
p("The barcode feature uses the browser's media APIs and JavaScript to read codes through the device camera, while server-side ZXing services generate EAN-13/QR assets and QuestPDF creates label documents. Therefore, ordinary phones, tablets, and computer cameras can scan products without a separate physical scanner or native mobile application. After a POS sale is completed, the interface presents a Sales Transaction Summary containing the finalized transaction details.")
p("Email is integrated through MailKit and MimeKit using configured SMTP credentials and STARTTLS where supported. It serves account confirmation/recovery and document delivery. Because checked-in SMTP fields are placeholders, its production operation depends on valid secret configuration. SQL Server integration is handled by Entity Framework Core and migrations. The chatbot uses Python FastAPI, Google Gemini, ODBC Driver 18, and a SQL connection intended to be read-only for assistant reporting.")

h("3. Additional Technical Details",1)
h("3.1 Data Flow",2)
flows=[
("Sales transaction","Employee scans/selects product -> POS builds cart -> server reloads authoritative products -> serializable transaction validates stock -> stock is deducted -> Transaction and TransactionItem records are stored -> transaction commits -> Sales Transaction Summary is displayed -> safety-stock metrics are recalculated."),
("Appointment and parts","Customer submits booking -> server derives customer identity and checks conflicts -> staff confirms, assigns mechanic, and selects products -> quantities may be reserved -> completion service creates linked sale and deducts stock atomically -> appointment becomes completed -> safety stock is recalculated."),
("Motorcycle build","Customer chooses motorcycle/package/parts -> compatibility is checked -> build request is stored -> employee reviews and may link installation appointment -> checkout validates selected parts -> linked sale and stock deduction commit -> build status becomes completed."),
("Procurement","Demand and completed-order history feed safety-stock calculations -> inventory position is compared with reorder policy -> eligible quantities are rounded and grouped by supplier -> a draft order slip is generated -> staff reviews and confirms it -> the confirmed order slip can be emailed to the supplier -> receiving updates incoming/current quantities and lead-time history."),
("Reports","The Dashboard and Transaction History pages retrieve current records from the StockSense database and summarize them into inventory indicators, recent sales, appointments, motorcycle builds, order-slip status, and transaction histories. When requested, the system can also generate PDF documents for confirmed order slips and barcode labels. The displayed information reflects the latest committed database records at the time the page is loaded or refreshed."),
("AI assistance","Browser sends question to authenticated ASP.NET endpoint -> server derives role -> FastAPI validates message/history -> permitted read tool retrieves current facts when needed -> Gemini synthesizes answer -> reply returns to the assistance component.")]
table(["Process","Data flow"],flows,[1.45,5.05])

h("3.2 Security Implementation",2)
h("Authentication",3)
p("ASP.NET Core Identity manages confirmed accounts and cookie authentication. The cookie has a 15-minute sliding lifetime, and repeated failed attempts trigger a short lockout. Login-specific and global IP rate limits reduce automated abuse. Outside development, the pipeline enables HSTS and HTTPS redirection. Deployment should set secure cookies to Always behind a correctly configured HTTPS proxy rather than relying only on SameAsRequest.")
h("Authorization",3)
p("Role claims and policies protect many staff and administrator operations, and server-side customer identity prevents one customer from selecting another customer's record ownership. The AI proxy also derives role from claims. Before production, a deny-by-default authorization review should correct the identified mechanics, services, pre-built, and product-data exposures and add tests proving Customers cannot call staff mutations.")
h("Password Protection",3)
p("Passwords are not stored as readable text. The project uses a SHA-256 pre-hash followed by BCrypt with work factor 12, integrated with Identity. Confirmation, recovery, lockout, and two-factor management screens complement hashing. Secrets and passwords must never be placed in source-controlled settings; the hard-coded fallback development SQL credential in Program.cs should be removed.")
h("Input and Data Protection",3)
p("DTO data annotations validate several important lengths, ranges, required values, and identifiers. EF Core parameterizes ordinary queries. High-risk inventory workflows use transactions and concurrency tokens. Product images are validated through server-side logic. Validation is not yet uniform across every DTO, so endpoint-specific validation and output filtering should be expanded. HTML email content should be encoded before interpolation.")
h("API and Integration Security",3)
p("The ASP.NET pipeline includes rate limiting, authentication, authorization, and antiforgery services/middleware. However, cookie-authenticated MVC mutations do not uniformly show antiforgery validation, so full CSRF protection should not be claimed until a global filter or explicit validation is verified. The FastAPI chatbot endpoint has strict schema and size limits but no direct service authentication and trusts its JSON role. It should accept a signed application token or reside on a private network. Its database user should have SELECT-only access, and the legacy staff arbitrary-SELECT path should be replaced by fixed, parameterized reports.")

h("3.3 Integration Points",2)
integrations=[
("Database","Entity Framework Core uses SQL Server for identity and operational records. Migrations define and evolve the schema; retry behavior and startup migration checks are configured. LocalDB supports development, while production requires a protected SQL Server/Azure SQL connection string and backup plan."),
("AI service","The .NET assistance controller proxies authenticated requests to FastAPI. FastAPI uses permitted tools and Google Gemini to compose answers. Timeout and base URL are configurable. Production requires private networking or service authentication and a least-privilege database identity."),
("Email service","MailKit/MimeKit connects to SMTP for account and document messages. Credentials, sender, host, port, and encryption settings must be injected securely. Failures should be logged without exposing credentials or message contents."),
("Documents and barcode","QuestPDF generates PDFs; ZXing creates barcode/QR content; browser JavaScript scans through a camera. Short-lived unpredictable download tokens are used for generated PDFs, although authorization on the download route should be reviewed."),
("External systems not present","No payment gateway, SMS provider, accounting/ERP connector, supplier purchasing API, or general BI platform is evidenced. These must be presented as future integrations rather than current capabilities.")]
table(["Integration","Technical description and boundary"],integrations,[1.45,5.05])

h("Implementation Limitations and Recommended Hardening",2)
p("Before production acceptance, the team should enforce staff/admin roles on every management mutation, filter customer-facing product DTOs, add consistent CSRF validation for cookie-authenticated APIs, remove fallback credentials from source, encode generated email HTML, and protect the chatbot with signed service authentication or private networking. It should also use a dedicated read-only chatbot database account and replace flexible staff SQL with predefined reports. These corrections do not change the architecture; they close trust-boundary gaps revealed by implementation review.")
p("The final deployment plan should name the actual hosting service, SQL tier, region, backup schedule, monitoring/alerting platform, secret store, SMTP provider, supported Python version, and recovery objectives. Load, accessibility, mobile, concurrency, and security testing should be recorded as evidence. Until those choices and tests are complete, they should be described as requirements rather than finished deployment characteristics.")

h("Technical Basis of This Chapter",2)
p("The description was derived from the StockSense solution and project files, ASP.NET startup configuration, controllers, domain entities, Entity Framework context and migrations, infrastructure services, Blazor pages and components, automated tests, and the separate Gemini_Chatbot FastAPI source and dependency list as inspected on 5 August 2026. Package versions are those declared by the project files. Recommended hardware and network values are planning baselines and should be replaced with measured institutional or production specifications when available.")

doc.save(OUT)
print(OUT)
